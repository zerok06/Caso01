from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import re
import os
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle, Image
from reportlab.lib.units import inch
from reportlab.lib import colors
from typing import Dict, Any
from datetime import datetime
from fastapi.responses import FileResponse
from fastapi import HTTPException, status
import tempfile
from typing import Dict, Any
from io import BytesIO
import logging

logger = logging.getLogger(__name__)

def generate_document(proposal_data: Dict[str, Any], format: str = "docx") -> bytes:
    """
    Genera un documento con la propuesta comercial.
    
    Args:
        proposal_data: Datos del análisis de la propuesta
        format: Formato del documento ("docx" o "pdf")
        
    Returns:
        Documento generado como bytes
        
    Raises:
        HTTPException 400: Si el formato no es soportado
        HTTPException 500: Si hay error en la generación
    """
    try:
        if format.lower() == "docx":
            return _generate_docx(proposal_data)
        elif format.lower() == "pdf":
            return _generate_pdf(proposal_data)
        else:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Formato no soportado: {format}. Use 'docx' o 'pdf'."
            )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error al generar documento ({format}): {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error al generar el documento: {str(e)}"
        )

def _generate_docx(proposal_data: Dict[str, Any]) -> bytes:
    """
    Genera un docuemto de Word basado en los datos de la propuesta proporcionados.

    Args:
        proposal_data (Dict[str, Any]): La información puede incluir título, descripción, objetivos, cronograma, presupuesto, etc.

    Returns:
        str: La ruta del archivo al documento de Word generado.
    """
    # Crear documento Word
    doc = Document()

    # ============ PORTADA IMAGEN ============
    # Determinar ruta absoluta de la imagen basándonos en la ubicación de este script
    current_dir = os.path.dirname(os.path.abspath(__file__)) # .../backend/core
    backend_dir = os.path.dirname(current_dir)               # .../backend
    image_path = os.path.join(backend_dir, 'src', 'assets', 'portada_tivit.jpg')
    
    if os.path.exists(image_path):
        try:
            # Añadir imagen centrada
            doc.add_picture(image_path, width=Inches(6))
            last_paragraph = doc.paragraphs[-1] 
            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Salto de página
            doc.add_page_break()
        except Exception as e:
            logger.error(f"Error al adjuntar imagen de portada en _generate_docx: {e}")
    else:
        logger.warning(f"No se encontró imagen de portada en: {image_path}")

    # ============ CONFIGURAR TIPOGRAFÍA GLOBAL ============
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(50, 50, 50)  # Gris oscuro elegante

    # ============ ESTILOS PERSONALIZADOS ============

    # Título Corporativo
    title_style = doc.styles.add_style('CorporateTitle', 1)
    title_style.font.name = 'Calibri Light'
    title_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri Light')
    title_style.font.size = Pt(32)
    title_style.font.color.rgb = RGBColor(30, 30, 30)

    # Encabezados de sección
    header_style = doc.styles.add_style('CorporateHeader', 1)
    header_style.font.name = 'Segoe UI Semibold'
    header_style.font.size = Pt(18)
    header_style.font.color.rgb = RGBColor(40, 40, 120)  # Azul corporativo

    # Encabezados de subsección
    subheader_style = doc.styles.add_style('CorporateSubHeader', 1)
    subheader_style.font.name = 'Segoe UI'
    subheader_style.font.size = Pt(14)
    subheader_style.font.color.rgb = RGBColor(70, 70, 70)

    # ============ PORTADA ============
    title = doc.add_paragraph("PROPUESTA TÉCNICA", style="CorporateTitle")
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("\n")
    doc.add_paragraph(f"Fecha: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph(f"Cliente: {proposal_data.get('cliente', 'No especificado')}")
    doc.add_paragraph("\nDocumento confidencial — Uso exclusivo del cliente.")

    doc.add_page_break()

    # ============ SECCIÓN: INTRODUCCIÓN ============
    doc.add_paragraph("1. Introducción", style="CorporateHeader")
    doc.add_paragraph(
        "El presente documento presenta la propuesta técnica elaborada por nuestro equipo, "
        "definida con base en los lineamientos y requerimientos compartidos por el cliente."
    )

    # ============ SECCIÓN: OBJETIVO ============
    doc.add_paragraph("2. Objetivo del Proyecto", style="CorporateHeader")
    
    # Manejar objetivo_general (puede ser lista) u objetivo_proyecto (string)
    objetivo = proposal_data.get("objetivo_general") or proposal_data.get("objetivo_proyecto")
    
    if objetivo:
        if isinstance(objetivo, list):
            for obj in objetivo:
                doc.add_paragraph(f"• {obj}", style="List Bullet")
        else:
            doc.add_paragraph(str(objetivo))
    else:
        doc.add_paragraph("Presentar una solución técnica integral que responda a las necesidades identificadas.")

    # ============ SECCIÓN: INFORMACIÓN GENERAL ============
    doc.add_paragraph("3. Información General del Proyecto", style="CorporateHeader")
    doc.add_paragraph(f"Fecha estimada de entrega: {proposal_data.get('fecha_entrega', 'No especificada')}")
    doc.add_paragraph(f"Cliente: {proposal_data.get('cliente', 'No especificado')}")

    # ============ SECCIÓN: ALCANCE ECONÓMICO ============
    doc.add_paragraph("4. Alcance Económico", style="CorporateHeader")
    alcance = proposal_data.get('alcance_economico', {})
    doc.add_paragraph(f"Presupuesto Estimado: {alcance.get('presupuesto', 'No especificado')}")
    doc.add_paragraph(f"Moneda: {alcance.get('moneda', 'No especificada')}")

    doc.add_paragraph(
        "Los valores expresados están sujetos a validación final y pueden ajustarse de acuerdo "
        "con requerimientos adicionales o cambios en el alcance del proyecto."
    )

    # ============ SECCIÓN: TECNOLOGÍAS ============
    doc.add_paragraph("5. Tecnologías Propuestas", style="CorporateHeader")

    tecnologias = proposal_data.get('tecnologias_requeridas', [])
    if tecnologias:
        p = doc.add_paragraph("Las tecnologías sugeridas son las siguientes:")
        for tech in tecnologias:
            doc.add_paragraph(f"• {tech}", style="List Bullet")
    else:
        doc.add_paragraph("No se especificaron tecnologías requeridas.")

    # ============ SECCIÓN: PREGUNTAS ============
    doc.add_paragraph("6. Preguntas para Aclaración", style="CorporateHeader")
    preguntas = proposal_data.get('preguntas_sugeridas', [])

    if preguntas:
        for pregunta in preguntas:
            doc.add_paragraph(f"• {pregunta}", style="List Bullet")
    else:
        doc.add_paragraph("No se registraron preguntas.")

    # ============ SECCIÓN: EQUIPO ============
    doc.add_paragraph("7. Equipo Propuesto", style="CorporateHeader")
    equipo = proposal_data.get('equipo_sugerido', [])

    for miembro in equipo:
        doc.add_paragraph(miembro.get('nombre', 'Sin nombre'), style="CorporateSubHeader")
        doc.add_paragraph(f"Rol: {miembro.get('rol', 'No especificado')}")
        doc.add_paragraph(f"Experiencia: {miembro.get('experiencia', 'No especificada')}")

        skills = miembro.get('skills', [])
        if skills:
            doc.add_paragraph("Habilidades Clave:")
            for skill in skills:
                doc.add_paragraph(f"• {skill}", style="List Bullet")
    
    # Guardar documento temporalmente
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
        tmp_path = tmp_file.name
        doc.save(tmp_path)
        
    file_name = f"Propuesta_{proposal_data.get('cliente', 'Cliente')}.docx"
    
    # Guardar en memoria
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    # Devolver documento
    return output.getvalue()

def _generate_pdf(proposal_data: Dict[str, Any]) -> bytes:
    """Genera documento PDF usando ReportLab"""
    
    output = BytesIO()
    doc = SimpleDocTemplate(output, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
    
    styles = getSampleStyleSheet()
    story = []
    
    # ============ PORTADA IMAGEN ============
    current_dir = os.path.dirname(os.path.abspath(__file__))
    backend_dir = os.path.dirname(current_dir)
    image_path = os.path.join(backend_dir, 'src', 'assets', 'portada_tivit.jpg')
    
    if os.path.exists(image_path):
        try:
            # 6 pulgadas de ancho, alto ajustado proporcionalmente
            img = Image(image_path, width=6*inch, height=4*inch, kind='proportional')
            story.append(img)
            story.append(Spacer(1, 0.5*inch))
            story.append(PageBreak())
        except Exception as e:
            logger.error(f"Error adjuntando imagen PDF: {e}")
            
    # Estilos personalizados
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#1e40af'),
        spaceAfter=30,
        alignment=1  # Center
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.HexColor('#1e3a8a'),
        spaceAfter=12,
        spaceBefore=12
    )
    
    # Título
    story.append(Paragraph("Propuesta Comercial", title_style))
    story.append(Spacer(1, 0.3*inch))
    
    # Cliente
    story.append(Paragraph("Cliente", heading_style))
    story.append(Paragraph(proposal_data.get("cliente", "N/A"), styles['Normal']))
    story.append(Spacer(1, 0.15*inch))
    
    # Alcance Económico
    story.append(Paragraph("Alcance Económico", heading_style))
    alcance = proposal_data.get("alcance_economico", {})
    presupuesto_text = f"<b>Presupuesto:</b> {alcance.get('presupuesto', 'N/A')} {alcance.get('moneda', 'USD')}"
    story.append(Paragraph(presupuesto_text, styles['Normal']))
    story.append(Spacer(1, 0.15*inch))
    
    # Fecha Entrega
    story.append(Paragraph("Fecha de Entrega", heading_style))
    story.append(Paragraph(proposal_data.get("fecha_entrega", "N/A"), styles['Normal']))
    story.append(Spacer(1, 0.15*inch))
    
    #Objetivo General
    story.append(Paragraph("Objetivo General", heading_style))
    objetivo_general = proposal_data.get("objetivo_general", [])
    if isinstance(objetivo_general, list):
        for objetivo in objetivo_general:
            story.append(Paragraph(f"• {objetivo}", styles['Normal']))
    else:
        story.append(Paragraph(str(objetivo_general) if objetivo_general else "N/A", styles['Normal']))
    story.append(Spacer(1, 0.15*inch))
    
    # Preguntas Sugeridas
    preguntas = proposal_data.get("preguntas_sugeridas", [])
    if preguntas:
        story.append(Paragraph("Preguntas Sugeridas", heading_style))
        for pregunta in preguntas:
            story.append(Paragraph(f"• {pregunta}", styles['Normal']))
        story.append(Spacer(1, 0.15*inch))
    
    # Equipo Sugerido
    equipo = proposal_data.get("equipo_sugerido", [])
    if equipo:
        story.append(Paragraph("Equipo Sugerido", heading_style))
        for miembro in equipo:
            nombre = miembro.get('nombre', 'N/A')
            rol = miembro.get('rol', 'N/A')
            exp = miembro.get('experiencia', 'N/A')
            skills = ', '.join(miembro.get('skills', []))
            
            miembro_text = f"<b>{nombre}</b><br/>Rol: {rol}<br/>Experiencia: {exp}<br/>Skills: {skills}"
            story.append(Paragraph(miembro_text, styles['Normal']))
            story.append(Spacer(1, 0.1*inch))
        story.append(Spacer(1, 0.15*inch))
    
    # Tecnologías
    techs = proposal_data.get("tecnologias_requeridas", [])
    if techs:
        story.append(Paragraph("Tecnologías Requeridas", heading_style))
        story.append(Paragraph(", ".join(techs), styles['Normal']))
    
    # Construir PDF
    doc.build(story)
    output.seek(0)
    
    # Devolver documento
    return output.getvalue()


# Dentro de document_service.py (donde tienes generate_document)

def generate_from_text(raw_text: str, format: str = "docx") -> bytes:
    """
    Genera un documento Word o PDF a partir de una cadena de texto formateada en Markdown.
    
    Args:
        raw_text: El texto de la propuesta formateado en Markdown.
        format: Formato del documento ("docx" o "pdf")
        
    Returns:
        Documento generado como bytes
        
    Raises:
        HTTPException 400: Si el formato no es soportado
        HTTPException 500: Si hay error en la generación
    """
    try:
        if format.lower() == "docx":
            return _generate_docx_from_markdown(raw_text)
        elif format.lower() == "pdf":
            return _generate_pdf_from_markdown(raw_text)
        else:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Formato no soportado: {format}. Use 'docx' o 'pdf'."
            )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error al generar documento desde Markdown ({format}): {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error al generar el documento: {str(e)}"
        )
        
# Funciones Helper
def _parse_markdown_table_data(lines):
    """Extrae las filas y columnas de una tabla Markdown."""
    table_data = []
    
    # 1. Ignorar la línea de alineación (ej: |:---|:---|:---|)
    if len(lines) > 1 and all(c in '-|: ' for c in lines[1]):
        lines = [lines[0]] + lines[2:]

    for line in lines:
        # Remover el primer y último pipe, dividir por pipe y limpiar espacios
        cells = [cell.strip() for cell in line.strip().strip('|').split('|')]
        if any(cells): # Solo añadir si la fila tiene contenido
            table_data.append(cells)
            
    return table_data

def _add_table_to_docx(doc: Document, table_data: list):
    """Añade una tabla formateada a un documento DOCX."""
    if not table_data:
        return

    num_rows = len(table_data)
    num_cols = len(table_data[0]) if table_data else 0
    
    if num_cols == 0:
        return

    # Crea la tabla con el número correcto de filas y columnas
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid' # Estilo básico de tabla

    for r_idx, row_data in enumerate(table_data):
        # Asegurarse de que la fila tenga el número correcto de columnas para evitar IndexError
        row_data_padded = row_data + [''] * (num_cols - len(row_data))
        
        row = table.rows[r_idx]
        for c_idx, cell_text in enumerate(row_data_padded):
            cell = row.cells[c_idx]
            cell.text = cell_text
            
            # Aplicar negrita a la fila de encabezado
            if r_idx == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        
                        
def _generate_docx_from_markdown(raw_text: str) -> bytes:
    """
    Genera un documento de Word basado en texto Markdown.
    """
    doc = Document()
    
    # ============ PORTADA ============
    # Determinar ruta absoluta de la imagen basándonos en la ubicación de este script
    current_dir = os.path.dirname(os.path.abspath(__file__)) # .../backend/core
    backend_dir = os.path.dirname(current_dir)               # .../backend
    image_path = os.path.join(backend_dir, 'src', 'assets', 'portada_tivit.jpg')
    
    if os.path.exists(image_path):
        try:
            # Añadir imagen centrada
            # 6 pulgadas de ancho es un buen estándar para portadas A4/Letter
            doc.add_picture(image_path, width=Inches(6))
            last_paragraph = doc.paragraphs[-1] 
            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Salto de página para empezar el contenido en la siguiente
            doc.add_page_break()
        except Exception as e:
            logger.error(f"Error al adjuntar imagen de portada: {e}")
            # Si falla, continuamos sin portada pero logueamos el error
    else:
        logger.warning(f"No se encontró imagen de portada en: {image_path}")

    # ============ CONFIGURAR ESTILOS ============
    # Puedes reusar los estilos que ya definiste (_generate_docx_from_json) o definirlos aquí
    # Para simplicidad, usaremos los estilos predeterminados (Heading 1, Heading 2, etc.)

    lines = raw_text.split('\n')
    story = [] # Lista para almacenar líneas de texto y tablas pendientes

    # Bucle de parseo
    for line in lines:
        line = line.strip()
        
        if not line:
            # Si story no está vacía y la última línea no era un párrafo, añade un espacio
            if story and isinstance(story[-1], tuple) and story[-1][1] != "":
                story.append(('P', '')) # Añadir un separador de párrafo

            continue

        # 1. Encabezados
        if line.startswith('# '):
            story.append(('H1', line.lstrip('# ').strip()))
        elif line.startswith('## '):
            story.append(('H2', line.lstrip('## ').strip()))
        elif line.startswith('### '):
            story.append(('H3', line.lstrip('### ').strip()))
        # 2. Tablas - Detecta el inicio de la sintaxis de tabla
        elif line.startswith('|') and line.count('|') > 2:
            story.append(('TABLE_START', line))
        # 3. Listas
        elif line.startswith('- ') or line.startswith('* '):
            story.append(('LIST', line.lstrip('-* ').strip()))
        # 4. Párrafos
        else:
            # Simplemente añade el texto. Si es la portada o resumen, se añadirá como texto normal.
            story.append(('P', line)) 
            
    # Lógica para construir el documento
    is_in_table = False
    current_table_lines = []

    for item in story:
        item_type, content = item
        
        # Manejo de Tablas
        if item_type == 'TABLE_START':
            is_in_table = True
            current_table_lines.append(content)
            continue
        
        # Si estábamos en una tabla y llega algo que NO es tabla, cerramos la tabla
        if is_in_table:
            # La tabla ha terminado, procesarla y reiniciar el estado
            table_data = _parse_markdown_table_data(current_table_lines)
            _add_table_to_docx(doc, table_data)
            
            is_in_table = False
            current_table_lines = []
        
        # Procesar items normales
        if item_type == 'H1':
            doc.add_heading(content, level=1)
        elif item_type == 'H2':
            doc.add_heading(content, level=2)
        elif item_type == 'H3':
            doc.add_heading(content, level=3)
        elif item_type == 'LIST':
            doc.add_paragraph(content, style='List Bullet')
        elif item_type == 'P':
            # Puedes aplicar negritas de Markdown (**texto**) con un parser más avanzado,
            # pero para esta versión simple, solo añade el párrafo.
            doc.add_paragraph(content)

    # Procesar la última tabla si el documento termina en una
    if is_in_table and current_table_lines:
        table_data = _parse_markdown_table_data(current_table_lines)
        _add_table_to_docx(doc, table_data)


    # Guardar en memoria y devolver bytes
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    return output.getvalue()


def _generate_pdf_from_markdown(raw_text: str) -> bytes:
    """
    Genera un documento PDF a partir de texto Markdown usando ReportLab (Implementación básica).
    """
    
    output = BytesIO()
    doc = SimpleDocTemplate(output, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
    
    styles = getSampleStyleSheet()
    story = []
    
    # ============ PORTADA IMAGEN ============
    current_dir = os.path.dirname(os.path.abspath(__file__))
    backend_dir = os.path.dirname(current_dir)
    image_path = os.path.join(backend_dir, 'src', 'assets', 'portada_tivit.jpg')
    
    if os.path.exists(image_path):
        try:
            img = Image(image_path, width=6*inch, height=4*inch, kind='proportional')
            story.append(img)
            story.append(Spacer(1, 0.5*inch))
            story.append(PageBreak())
        except Exception as e:
            logger.error(f"Error adjuntando imagen PDF markdown: {e}")
    
    # Estilos básicos de ReportLab para Markdown
    h1_style = styles['Heading1']
    h2_style = styles['Heading2']
    h3_style = styles['Heading3']
    normal_style = styles['Normal']
    
    lines = raw_text.split('\n')
    
    # Lógica de parseo simple para ReportLab
    for line in lines:
        line = line.strip()
        if not line:
            story.append(Spacer(1, 0.1 * inch))
            continue

        if line.startswith('# '):
            story.append(Paragraph(line.lstrip('# ').strip(), h1_style))
        elif line.startswith('## '):
            story.append(Paragraph(line.lstrip('## ').strip(), h2_style))
        elif line.startswith('### '):
            story.append(Paragraph(line.lstrip('### ').strip(), h3_style))
        elif line.startswith('- ') or line.startswith('* '):
            # Formato de lista simple
            list_item = line.lstrip('-* ').strip()
            story.append(Paragraph(f"• {list_item}", normal_style))
        elif line.startswith('|') and line.count('|') > 2:
            # Para tablas en PDF, se necesita un parser más sofisticado. 
            # Aquí solo incluiremos el texto de la tabla como un párrafo simple para no fallar.
            story.append(Paragraph(f"**TABLA (Formato ReportLab no implementado):** {line}", normal_style))
        else:
            # Párrafos normales
            story.append(Paragraph(line, normal_style))

    # Construir PDF
    doc.build(story)
    output.seek(0)
    
    return output.getvalue()