import os
from pathlib import Path
import pandas as pd
import pypdf
import docx
import openpyxl 
import mimetypes
from typing import Generator

def extract_text_from_file(file_path: Path) -> Generator[str, None, None]:
    """
    Extrae texto de un archivo (PDF, DOCX, XLSX, CSV, TXT) de manera eficiente (streaming).
    Devuelve un generador que emite chunks de texto.
    """
    # Intentar detectar tipo MIME
    file_type = mimetypes.guess_type(file_path)[0]
    
    # Fallback: detectar por extensión si mimetypes falla
    file_extension = Path(file_path).suffix.lower()
    
    if not file_type:
        extension_to_mimetype = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            '.csv': 'text/csv',
            '.txt': 'text/plain',
        }
        file_type = extension_to_mimetype.get(file_extension)
        print(f"PARSER: MIME type no detectado, usando extensión: {file_extension} -> {file_type}")
    
    print(f"PARSER: Iniciando extracción streaming de {file_path} (Tipo: {file_type})")
    
    try:
        if file_type == "application/pdf" or file_extension == '.pdf':
            # Usar pypdf (reemplazo moderno de PyPDF2)
            try:
                reader = pypdf.PdfReader(file_path)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        yield text + "\n"
            except Exception as pdf_error:
                print(f"PARSER: Error con pypdf, intentando pdfplumber: {pdf_error}")
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            yield text + "\n"
        
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or file_extension == '.docx':
            doc = docx.Document(file_path)
            # Yield por párrafos para no cargar todo
            current_chunk = []
            for para in doc.paragraphs:
                current_chunk.append(para.text)
                if len(current_chunk) >= 10: # Agrupar cada 10 párrafos
                    yield "\n".join(current_chunk) + "\n"
                    current_chunk = []
            
            if current_chunk:
                yield "\n".join(current_chunk) + "\n"

            # Tablas
            for table in doc.tables:
                for row in table.rows:
                    row_text = " ".join(cell.text for cell in row.cells)
                    yield row_text + "\n"
        
        elif file_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet" or file_extension == '.xlsx':
            # Usar openpyxl en modo read_only para streaming real
            wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            for sheet in wb.worksheets:
                yield f"--- Hoja: {sheet.title} ---\n"
                for row in sheet.iter_rows(values_only=True):
                    # Filtrar Nones y convertir a string
                    row_text = ", ".join(str(cell) for cell in row if cell is not None)
                    if row_text:
                        yield row_text + "\n"
        
        elif file_type == "text/csv" or file_extension == '.csv':
            # Usar pandas con chunksize
            for chunk in pd.read_csv(file_path, chunksize=1000):
                yield chunk.to_string() + "\n"
            
        elif file_type == "text/plain" or file_extension == '.txt':
            with open(file_path, "r", encoding="utf-8") as f:
                while True:
                    chunk = f.read(4096) # Leer 4KB a la vez
                    if not chunk:
                        break
                    yield chunk
        
        else:
            print(f"PARSER: Tipo de archivo '{file_type}' no soportado.")
            raise ValueError(f"Tipo de archivo no soportado: {file_type}")
        
        print(f"PARSER: Extracción finalizada para {file_path}")

    except Exception as e:
        print(f"PARSER: Error al extraer texto de {file_path}: {e}")
        raise